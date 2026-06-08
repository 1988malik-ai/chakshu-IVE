"""Automated reports — HTML, PDF, DOCX with processing steps, settings, and references."""

from __future__ import annotations

import html
import json
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any

from aive.project.workflow import AiveProject
from aive.reports.context import ReportContext, build_context

from aive.i18n.translations import Translator


@dataclass
class ReportSettings:
    paper_size: str = "A4"
    orientation: str = "portrait"
    template: str = "standard"
    title: str = "Chakshu Processing Report"
    author: str = ""
    locale: str = "en"
    include_steps: bool = True
    include_settings: bool = True
    include_references: bool = True
    include_notes: bool = True
    include_bookmarks: bool = True
    include_pipeline: bool = True
    output_formats: list[str] = field(default_factory=lambda: ["html", "pdf", "docx"])


TEMPLATES = {
    "standard": {"accent": "#1e40af", "header_bg": "#f1f5f9"},
    "detailed": {"accent": "#0f766e", "header_bg": "#ecfdf5"},
    "executive": {"accent": "#4c1d95", "header_bg": "#f5f3ff"},
    "minimal": {"accent": "#374151", "header_bg": "#ffffff"},
}


def generate_report(
    project: AiveProject,
    output_dir: Path,
    settings: ReportSettings | None = None,
    *,
    case_meta: dict[str, Any] | None = None,
    bookmarks: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    settings = settings or ReportSettings()
    tr = Translator(settings.locale)
    if settings.title in ("Chakshu Processing Report", "", "AI-IVE Processing Report"):
        settings.title = tr.tr("report.title")

    output_dir.mkdir(parents=True, exist_ok=True)
    ctx = build_context(project, case=case_meta, bookmarks=bookmarks)
    tpl = TEMPLATES.get(settings.template, TEMPLATES["standard"])
    stem = f"report_{project.project_id[:8]}"
    outputs: list[dict[str, Any]] = []
    errors: list[str] = []

    formats = {f.lower().replace("doc", "docx") for f in settings.output_formats}

    if "html" in formats:
        path = output_dir / f"{stem}.html"
        path.write_text(_render_html(ctx, settings, tpl, tr), encoding="utf-8")
        outputs.append({"format": "html", "path": str(path)})

    if "pdf" in formats:
        path = output_dir / f"{stem}.pdf"
        result = _render_pdf(ctx, settings, tpl, tr, path)
        if result.get("success"):
            outputs.append({"format": "pdf", "path": str(path)})
        else:
            errors.append(result.get("error", "pdf failed"))

    if "docx" in formats:
        path = output_dir / f"{stem}.docx"
        result = _render_docx(ctx, settings, tr, path)
        if result.get("success"):
            outputs.append({"format": "docx", "path": str(path)})
        else:
            rtf = output_dir / f"{stem}.rtf"
            rtf.write_text(_render_rtf(ctx, settings), encoding="utf-8")
            outputs.append({"format": "rtf", "path": str(rtf), "note": result.get("error", "docx_unavailable")})
            errors.append(result.get("error", "python-docx not installed"))

    return {
        "success": bool(outputs),
        "outputs": outputs,
        "errors": errors,
        "paper_size": settings.paper_size,
        "step_count": len(project.workflow_steps),
        "formats_requested": list(formats),
    }


def _meta_block(ctx: ReportContext, settings: ReportSettings, tr: Translator) -> list[tuple[str, str]]:
    p = ctx.project
    rows = [
        (tr.tr("report.project"), p.name),
        (tr.tr("report.id"), p.project_id),
        (tr.tr("report.paper"), f"{settings.paper_size} ({settings.orientation})"),
        (tr.tr("report.template"), settings.template),
        (tr.tr("report.generated"), f"{datetime.utcnow().isoformat()}Z"),
        (tr.tr("report.steps_count"), str(len(p.workflow_steps))),
    ]
    if settings.author:
        rows.append((tr.tr("report.author"), settings.author))
    if ctx.case:
        if ctx.case.get("case_number") or ctx.case.get("display_id"):
            rows.append((tr.tr("report.case_number"), ctx.case.get("display_id") or ctx.case.get("case_number", "")))
        if ctx.case.get("examiner"):
            rows.append((tr.tr("report.examiner"), ctx.case["examiner"]))
        if ctx.case.get("agency"):
            rows.append((tr.tr("report.agency"), ctx.case["agency"]))
    return rows


def _render_html(ctx: ReportContext, settings: ReportSettings, tpl: dict, tr: Translator) -> str:
    meta_html = "".join(f"<p><strong>{html.escape(k)}:</strong> {html.escape(v)}</p>" for k, v in _meta_block(ctx, settings, tr))

    steps_html = ""
    if settings.include_steps:
        for row in ctx.step_rows(
            include_settings=settings.include_settings,
            include_references=settings.include_references,
        ):
            steps_html += f"""
        <tr>
          <td>{row['index']}</td>
          <td>{html.escape(row['timestamp'])}</td>
          <td><strong>{html.escape(row['action'])}</strong></td>
          <td><pre class="settings-cell">{html.escape(row['settings'])}</pre></td>
          <td>{html.escape(row['references'])}</td>
        </tr>"""

    notes_html = ""
    if settings.include_notes and ctx.project.examination_notes:
        for n in ctx.project.examination_notes:
            notes_html += f"<li><strong>{html.escape(str(n.get('created', '')))}</strong> — {html.escape(n.get('text', n.get('body', '')))}</li>"

    bookmarks_html = ""
    if settings.include_bookmarks and ctx.bookmarks:
        for b in ctx.bookmarks[:50]:
            bookmarks_html += (
                f"<li>{html.escape(b.get('label', b.get('id', '')))} "
                f"({html.escape(b.get('bookmark_type', b.get('type', '')))}) "
                f"@ {b.get('time_sec', '—')}s</li>"
            )

    pipeline_block = ""
    if settings.include_pipeline:
        pipeline_block = f"""
<h2>{tr.tr('report.filter_pipeline')}</h2>
<pre>{html.escape(ctx.pipeline_summary())}</pre>"""

    return f"""<!DOCTYPE html>
<html lang="{settings.locale}"><head><meta charset="utf-8"><title>{html.escape(settings.title)}</title>
<style>
  body {{ font-family: 'Segoe UI', system-ui, sans-serif; margin: 40px; color: #1e293b; line-height: 1.45; }}
  h1 {{ color: {tpl['accent']}; border-bottom: 3px solid {tpl['accent']}; padding-bottom: 8px; }}
  h2 {{ color: {tpl['accent']}; margin-top: 28px; font-size: 1.1rem; }}
  .meta {{ background: {tpl['header_bg']}; padding: 16px; border-radius: 8px; margin: 20px 0; }}
  table {{ width: 100%; border-collapse: collapse; margin-top: 12px; font-size: 13px; }}
  th {{ background: {tpl['accent']}; color: white; text-align: left; padding: 10px; }}
  td {{ border-bottom: 1px solid #e2e8f0; padding: 8px; vertical-align: top; }}
  .settings-cell {{ margin: 0; font-size: 11px; white-space: pre-wrap; max-width: 320px; }}
  .footer {{ margin-top: 40px; font-size: 12px; color: #64748b; }}
  ul {{ padding-left: 20px; }}
</style></head><body>
<h1>{html.escape(settings.title)}</h1>
<div class="meta">{meta_html}</div>
<h2>{tr.tr('report.processing_steps')}</h2>
<table>
  <tr><th>#</th><th>{tr.tr('report.col.time')}</th><th>{tr.tr('report.col.action')}</th>
      <th>{tr.tr('report.col.settings')}</th><th>{tr.tr('report.col.references')}</th></tr>
  {steps_html or f'<tr><td colspan="5">{tr.tr("report.no_steps")}</td></tr>'}
</table>
{pipeline_block}
{f'<h2>{tr.tr("report.examination_notes")}</h2><ul>{notes_html}</ul>' if notes_html else ''}
{f'<h2>{tr.tr("report.bookmarks")}</h2><ul>{bookmarks_html}</ul>' if bookmarks_html else ''}
<div class="footer">{tr.tr('report.footer')}</div>
</body></html>"""


def _render_pdf(
    ctx: ReportContext,
    settings: ReportSettings,
    tpl: dict,
    tr: Translator,
    path: Path,
) -> dict[str, Any]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A3, A4, landscape, legal, letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError:
        return {"success": False, "error": "reportlab not installed — pip install reportlab"}

    sizes = {"A4": A4, "Letter": letter, "Legal": legal, "A3": A3}
    page = sizes.get(settings.paper_size, A4)
    if settings.orientation == "landscape":
        page = landscape(page)

    doc = SimpleDocTemplate(
        str(path),
        pagesize=page,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm,
    )
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9, leading=11)
    small = ParagraphStyle("Small", parent=styles["Normal"], fontSize=7, leading=9)
    story = [
        Paragraph(html.escape(settings.title), styles["Title"]),
        Spacer(1, 8),
    ]
    for label, value in _meta_block(ctx, settings, tr):
        story.append(Paragraph(f"<b>{html.escape(label)}:</b> {html.escape(value)}", body))
    story.append(Spacer(1, 12))
    story.append(Paragraph(tr.tr("report.processing_steps"), styles["Heading2"]))

    if settings.include_steps and ctx.project.workflow_steps:
        header = ["#", tr.tr("report.col.time"), tr.tr("report.col.action")]
        if settings.include_settings:
            header.append(tr.tr("report.col.settings"))
        if settings.include_references:
            header.append(tr.tr("report.col.references"))

        data = [header]
        for row in ctx.step_rows(
            include_settings=settings.include_settings,
            include_references=settings.include_references,
        ):
            line = [row["index"], row["timestamp"][:19], row["action"][:40]]
            if settings.include_settings:
                line.append(Paragraph(row["settings"][:500].replace("\n", "<br/>"), small))
            if settings.include_references:
                line.append(row["references"][:80])
            data.append(line)

        col_count = len(header)
        col_width = (doc.width - 10) / col_count
        table = Table(data, colWidths=[col_width] * col_count, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(tpl["accent"])),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                    ("FONTSIZE", (0, 0), (-1, 0), 9),
                    ("FONTSIZE", (0, 1), (-1, -1), 8),
                    ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        story.append(table)
    else:
        story.append(Paragraph(tr.tr("report.no_steps"), body))

    if settings.include_pipeline and ctx.project.filter_pipeline:
        story.append(Spacer(1, 12))
        story.append(Paragraph(tr.tr("report.filter_pipeline"), styles["Heading2"]))
        for line in ctx.pipeline_summary().split("\n")[:30]:
            story.append(Paragraph(html.escape(line), body))

    doc.build(story)
    return {"success": True}


def _render_docx(
    ctx: ReportContext,
    settings: ReportSettings,
    tr: Translator,
    path: Path,
) -> dict[str, Any]:
    try:
        from docx import Document
    except ImportError:
        return {"success": False, "error": "python-docx not installed — pip install python-docx"}

    doc = Document()
    doc.add_heading(settings.title, 0)
    for label, value in _meta_block(ctx, settings, tr):
        doc.add_paragraph(f"{label}: {value}")

    doc.add_heading(tr.tr("report.processing_steps"), level=1)
    cols = ["#", tr.tr("report.col.time"), tr.tr("report.col.action")]
    if settings.include_settings:
        cols.append(tr.tr("report.col.settings"))
    if settings.include_references:
        cols.append(tr.tr("report.col.references"))

    table = doc.add_table(rows=1, cols=len(cols))
    for i, h in enumerate(cols):
        table.rows[0].cells[i].text = h

    for row in ctx.step_rows(
        include_settings=settings.include_settings,
        include_references=settings.include_references,
    ):
        cells = table.add_row().cells
        cells[0].text = row["index"]
        cells[1].text = row["timestamp"]
        cells[2].text = row["action"]
        idx = 3
        if settings.include_settings:
            cells[idx].text = row["settings"]
            idx += 1
        if settings.include_references:
            cells[idx].text = row["references"]

    if not ctx.project.workflow_steps:
        doc.add_paragraph(tr.tr("report.no_steps"))

    if settings.include_pipeline and ctx.project.filter_pipeline:
        doc.add_heading(tr.tr("report.filter_pipeline"), level=1)
        doc.add_paragraph(ctx.pipeline_summary())

    if settings.include_notes and ctx.project.examination_notes:
        doc.add_heading(tr.tr("report.examination_notes"), level=1)
        for n in ctx.project.examination_notes:
            doc.add_paragraph(
                f"{n.get('created', '')}: {n.get('text', n.get('body', ''))}",
                style="List Bullet",
            )

    if settings.include_bookmarks and ctx.bookmarks:
        doc.add_heading(tr.tr("report.bookmarks"), level=1)
        for b in ctx.bookmarks[:50]:
            doc.add_paragraph(
                f"{b.get('label', '')} ({b.get('bookmark_type', '')}) — {b.get('time_sec', '')}s",
                style="List Bullet",
            )

    doc.save(str(path))
    return {"success": True}


def _render_rtf(ctx: ReportContext, settings: ReportSettings) -> str:
    lines = [r"{\rtf1\ansi", r"\b " + settings.title + r"\b0\par"]
    for row in ctx.step_rows():
        lines.append(
            rf"{row['index']}. [{row['timestamp']}] {row['action']} | {row['settings'][:120]} | {row['references']}\par"
        )
    lines.append("}")
    return "\n".join(lines)
